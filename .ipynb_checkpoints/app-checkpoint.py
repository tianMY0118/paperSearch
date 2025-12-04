"""
paperSearch - 智能学术文献助手 v1.0
Author: tianMY0118
Description: 从 arXiv 快速检索并导出学术论文，支持多格式输出。
License: MIT
"""

import gradio as gr
import requests
import feedparser
import sys
import io
import logging
from datetime import datetime
import os
import json

# ===== 第三方库 =====
from docx import Document
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# ===== 日志配置 =====
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [ScholarSift] %(levelname)s - %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)]
)
logger = logging.getLogger("ScholarSift")

# ===== 全局状态缓存 =====
paper_data_cache = []
current_query_info = {}

# 修复 stdout 编码（兼容 Windows）
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# ===== 核心功能 =====
def search_papers(query: str, max_results: int):
    global paper_data_cache, current_query_info
    logger.info(f"用户发起检索：关键词='{query}', 最多返回 {max_results} 篇")
    
    try:
        arxiv_url = f"http://export.arxiv.org/api/query?search_query=all:{query}&start=0&max_results={max_results}"
        response = requests.get(arxiv_url, timeout=15)
        response.raise_for_status()
    except Exception as e:
        logger.error(f"arXiv API 请求失败: {e}")
        return "❌ 无法连接 arXiv 服务，请检查网络或稍后重试。", None

    feed = feedparser.parse(response.text)
    if not feed.entries:
        logger.warning("未找到匹配论文")
        return "🔍 未找到相关论文。", None

    papers = []
    output = f"📚 ScholarSift 检索结果（关键词: {query}）\n\n"
    for idx, entry in enumerate(feed.entries):
        title = entry.title
        authors = ', '.join(author.name for author in entry.authors)
        published = entry.published.split("T")[0]
        abstract = entry.summary.replace('\n', ' ').strip()
        pdf_link = entry.links[1].href if len(entry.links) > 1 else entry.id

        papers.append({
            "Title": title,
            "Authors": authors,
            "Published": published,
            "PDF Link": pdf_link,
            "Abstract": abstract,
        })

        output += f"📄 论文 {idx + 1}\n"
        output += f"标题       : {title}\n"
        output += f"作者       : {authors}\n"
        output += f"发表日期   : {published}\n"
        output += f"PDF 链接   : {pdf_link}\n"
        output += f"摘要       : {abstract}\n"
        output += "—" * 60 + "\n\n"

    paper_data_cache = papers
    current_query_info = {
        "query": query,
        "max_results": max_results,
        "num_found": len(papers),
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }

    logger.info(f"检索完成，共找到 {len(papers)} 篇论文")
    return output, gr.update(visible=True)


def export_results(format: str):
    global paper_data_cache
    logger.info(f"用户请求导出为 {format} 格式")
    
    papers = paper_data_cache
    if not papers:
        logger.warning("导出失败：无缓存数据")
        return None

    os.makedirs("outputs", exist_ok=True)
    base_name = "scholarsift_export"
    ext_map = {"Text": "txt", "Word": "docx", "PDF": "pdf", "Excel": "xlsx"}
    filename = f"outputs/{base_name}.{ext_map.get(format, 'txt')}"

    watermark = "\n— 导出自 ScholarSift 智能学术助手 (https://yourwebsite.com) —\n"

    try:
        if format == "Word":
            doc = Document()
            doc.add_heading("ScholarSift 学术论文导出报告", 0)
            for i, p in enumerate(papers):
                doc.add_heading(f"论文 {i+1}: {p['Title']}", level=1)
                doc.add_paragraph(f"作者       : {p['Authors']}")
                doc.add_paragraph(f"发表日期   : {p['Published']}")
                doc.add_paragraph(f"PDF 链接   : {p['PDF Link']}")
                doc.add_paragraph(f"摘要       : {p['Abstract']}")
            doc.add_paragraph(watermark)
            doc.save(filename)

        elif format == "PDF":
            c = canvas.Canvas(filename, pagesize=letter)
            width, height = letter
            margin = 50
            y = height - margin

            def draw_line(text, size=10, spacing=14):
                nonlocal y
                c.setFont("Helvetica", size)
                for line in text.split('\n'):
                    if y < margin:
                        c.showPage()
                        y = height - margin
                        c.setFont("Helvetica", size)
                    c.drawString(margin, y, line[:100])  # 防止超宽
                    y -= spacing

            draw_line("ScholarSift 学术论文导出报告", size=14, spacing=20)
            y -= 10
            for i, p in enumerate(papers):
                draw_line(f"论文 {i+1}: {p['Title']}", size=12)
                draw_line(f"作者       : {p['Authors']}")
                draw_line(f"发表日期   : {p['Published']}")
                draw_line(f"PDF 链接   : {p['PDF Link']}")
                draw_line(f"摘要       : {p['Abstract']}")
                draw_line("—" * 70)
                y -= 10
            draw_line(watermark)
            c.save()

        elif format == "Excel":
            df = pd.DataFrame(papers)
            df.to_excel(filename, index=False)

        else:  # Text
            with open(filename, 'w', encoding='utf-8') as f:
                f.write("ScholarSift 学术论文导出报告\n\n")
                for i, p in enumerate(papers):
                    f.write(f"论文 {i+1}\n")
                    for k, v in p.items():
                        f.write(f"{k:<12}: {v}\n")
                    f.write("—" * 60 + "\n\n")
                f.write(watermark)

        # 记录导出日志（可选）
        log_event = {
            "tool": "ScholarSift",
            "action": "export",
            "format": format,
            "count": len(papers),
            "time": datetime.now().isoformat()
        }
        with open("scholarsift_export_log.json", "w", encoding="utf-8") as f:
            json.dump(log_event, f, ensure_ascii=False, indent=2)

        logger.info(f"导出成功: {filename}")
        return filename

    except Exception as e:
        logger.error(f"导出失败: {e}")
        return None


# ===== Gradio 界面 =====
with gr.Blocks(title="ScholarSift") as demo:
    gr.Markdown("## 🧠 ScholarSift — 智能学术文献助手")
    gr.Markdown("快速从 arXiv 检索前沿论文，并一键导出为 Word / PDF / Excel / 文本格式")

    with gr.Row():
        topic = gr.Textbox(
            label="研究主题关键词",
            placeholder="例如：大语言模型、量子计算、气候变化",
            value="Large Language Models"
        )
        num = gr.Slider(1, 10, value=5, step=1, label="最多返回篇数")

    search_btn = gr.Button("🔍 开始检索", variant="primary")
    result_box = gr.Textbox(label="检索结果", lines=20, max_lines=30)
    
    with gr.Row():
        export_dropdown = gr.Dropdown(
            choices=["Text", "Word", "PDF", "Excel"],
            label="选择导出格式",
            value="PDF",
            interactive=True
        )
        export_btn = gr.Button("📤 导出结果", variant="secondary")

    file_output = gr.File(label="📥 下载导出文件")

    # 事件绑定
    search_btn.click(
        fn=search_papers,
        inputs=[topic, num],
        outputs=[result_box, export_dropdown]
    )
    export_btn.click(
        fn=export_results,
        inputs=export_dropdown,
        outputs=file_output
    )

# ===== 启动服务 =====
if __name__ == "__main__":
    logger.info("启动 ScholarSift 服务...")
    demo.launch(
        server_name="0.0.0.0",
        server_port=7860,
        debug=True,
        favicon_path=None  # 可替换为你的 favicon.ico
    )